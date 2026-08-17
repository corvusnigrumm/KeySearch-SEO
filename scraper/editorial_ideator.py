"""
Módulo de Ideación Editorial, Extracción de Tags Reales de Google Trends
y Redacción Periodística Estructurada según el Manual de Referencia de Estilo.
"""

import io
import json
import logging

from config import GROQ_API_KEY
from scraper.ai_client import post_groq_json as _post_groq_json
from scraper.volume_estimator import HAS_PYTRENDS, TRENDS_TIMEFRAME

logger = logging.getLogger(__name__)

# Mapeo de códigos de país para pytrends
COUNTRY_GEO_MAP = {
    "Colombia": "CO",
    "España": "ES",
    "México": "MX",
    "Argentina": "AR",
    "Chile": "CL",
    "Perú": "PE",
    "Estados Unidos": "US",
    "Global": "",
}


def _normalizar_pais(pais: str) -> str:
    return COUNTRY_GEO_MAP.get(pais, "CO" if "colombia" in str(pais).lower() else "")


def obtener_tags_reales_google(
    keyword: str,
    sugerencias: list[str] = None,
    preguntas_paa: list[str] = None,
    pais: str = "Colombia",
    language_code: str = "es",
) -> dict:
    """
    Extrae etiquetas y taxonomías 100% REALES consultando:
    1. Google Trends (Rising Queries / Breakouts con aumentos +%)
    2. Google Trends (Top Related Queries)
    3. Google Trends (Related Topics / Entidades reconocidas)
    4. Google Autocomplete exacto (Long-tail)

    Garantía: Cero tags inventados o genéricos.
    """
    geo_code = _normalizar_pais(pais)
    sugs = [s.strip() for s in (sugerencias or []) if s and s.strip()]
    paas = [p.strip() for p in (preguntas_paa or []) if p and p.strip()]

    rising_queries = []
    top_queries = []
    related_topics = []

    if HAS_PYTRENDS and keyword:
        try:
            from pytrends.request import TrendReq

            pytrends = TrendReq(hl=language_code, tz=360, timeout=(10, 25))
            pytrends.build_payload(
                [keyword[:90]],
                cat=0,
                timeframe=TRENDS_TIMEFRAME,
                geo=geo_code,
                gprop="",
            )

            # 1. Consultas relacionadas (Top y Rising)
            try:
                related_dict = pytrends.related_queries()
                if keyword in related_dict and related_dict[keyword]:
                    data_kw = related_dict[keyword]

                    # Rising
                    if "rising" in data_kw and data_kw["rising"] is not None:
                        df_rising = data_kw["rising"]
                        if not df_rising.empty and "query" in df_rising.columns:
                            for _, row in df_rising.head(10).iterrows():
                                q_name = str(row["query"]).strip()
                                val = row.get("value", "")
                                val_str = "Breakout" if str(val).lower() == "breakout" or val > 1000 else f"+{val}%"
                                rising_queries.append(
                                    {"tag": q_name, "tipo": "Breakout / En Aumento", "crecimiento": val_str}
                                )

                    # Top
                    if "top" in data_kw and data_kw["top"] is not None:
                        df_top = data_kw["top"]
                        if not df_top.empty and "query" in df_top.columns:
                            for _, row in df_top.head(10).iterrows():
                                q_name = str(row["query"]).strip()
                                val = row.get("value", 100)
                                top_queries.append(
                                    {"tag": q_name, "tipo": "Top Búsqueda", "interes_relativo": int(val)}
                                )
            except Exception as e:
                logger.warning("Error extrayendo related_queries en pytrends: %s", e)

            # 2. Temas / Entidades relacionadas (Related Topics)
            try:
                topics_dict = pytrends.related_topics()
                if keyword in topics_dict and topics_dict[keyword]:
                    data_topics = topics_dict[keyword]
                    if "top" in data_topics and data_topics["top"] is not None:
                        df_top_t = data_topics["top"]
                        if not df_top_t.empty and "topic_title" in df_top_t.columns:
                            for _, row in df_top_t.head(8).iterrows():
                                t_title = str(row["topic_title"]).strip()
                                t_type = str(row.get("topic_type", "Entidad")).strip()
                                related_topics.append({"tag": t_title, "tipo": t_type})
            except Exception as e:
                logger.warning("Error extrayendo related_topics en pytrends: %s", e)

        except Exception as e:
            logger.warning("Error general en pytrends para tags: %s", e)

    # Autocompletado long-tail
    long_tail_exacto = [s for s in sugs if len(s.split()) >= 3][:12]
    if not long_tail_exacto and sugs:
        long_tail_exacto = sugs[:8]

    # Construir clasificación estricta de tags
    entidades_marcas = [t["tag"] for t in related_topics if t["tag"].lower() != keyword.lower()]
    tendencias_breakout = [q["tag"] for q in rising_queries]
    alto_volumen_top = [q["tag"] for q in top_queries if q["tag"].lower() != keyword.lower()]

    # Fallback si pytrends no devolvió nada (usar sugerencias reales)
    if not alto_volumen_top and sugs:
        alto_volumen_top = sugs[:6]
    if not tendencias_breakout and paas:
        tendencias_breakout = [p.replace("¿", "").replace("?", "").strip() for p in paas[:4]]

    # Lista plana depurada sin duplicados
    tags_todos = []
    vistos = set()

    def _agregar_tag(t):
        t_clean = str(t).strip()
        t_lower = t_clean.lower()
        if t_clean and t_lower not in vistos and len(t_clean) >= 2:
            vistos.add(t_lower)
            tags_todos.append(t_clean)

    # Prioridad: Keyword base -> Entidades -> Breakouts -> Top Queries -> Long Tail
    _agregar_tag(keyword)
    for t in entidades_marcas:
        _agregar_tag(t)
    for t in tendencias_breakout:
        _agregar_tag(t)
    for t in alto_volumen_top:
        _agregar_tag(t)
    for t in long_tail_exacto:
        _agregar_tag(t)

    return {
        "keyword_principal": keyword,
        "pais": pais,
        "tags_clasificados": {
            "entidades_y_marcas": entidades_marcas[:5],
            "tendencias_breakout": tendencias_breakout[:6],
            "alto_volumen_top": alto_volumen_top[:6],
            "long_tail_exacto": long_tail_exacto[:8],
        },
        "detalles_trends": {
            "rising": rising_queries,
            "top": top_queries,
            "topics": related_topics,
        },
        "tags_lista_plana": tags_todos[:15],
        "tags_csv_ready": ", ".join(tags_todos[:12]),
    }


def _detectar_categoria_contextual(keyword: str, sugerencias: list[str] = None, preguntas_paa: list[str] = None) -> str:
    """Clasifica la keyword en su dominio temático real para evitar titulares disparatados."""
    full_text = " ".join([str(keyword)] + (sugerencias or []) + (preguntas_paa or [])).lower()

    if any(
        k in full_text
        for k in [
            "terremoto",
            "sismo",
            "temblor",
            "epicentro",
            "magnitud",
            "replica",
            "sismologico",
            "desastre",
            "tsunami",
            "volcan",
            "inundacion",
            "lluvias",
            "huracan",
            "emergencia",
            "muertos",
            "heridos",
            "alerta",
        ]
    ):
        return "noticias_sismos"
    if any(
        k in full_text
        for k in [
            "banco",
            "tarjeta",
            "credito",
            "debito",
            "cuenta",
            "ahorros",
            "prestamo",
            "hipoteca",
            "interes",
            "subsidio",
            "impuesto",
            "dolar",
            "tasa",
            "nequi",
            "daviplata",
            "bancolombia",
            "davivienda",
            "bbva",
            "finanzas",
        ]
    ):
        return "finanzas_banca"
    if any(
        k in full_text
        for k in [
            "iphone",
            "samsung",
            "xiaomi",
            "huawei",
            "whatsapp",
            "android",
            "ios",
            "windows",
            "laptop",
            "procesador",
            "inteligencia artificial",
            "chatgpt",
            "app",
            "consola",
            "playstation",
        ]
    ):
        return "tecnologia"
    if any(
        k in full_text
        for k in [
            "receta",
            "cocinar",
            "ingredientes",
            "como preparar",
            "comida",
            "postre",
            "almuerzo",
            "cena",
            "plato",
            "gastronomia",
            "sabor",
            "cocina",
        ]
    ):
        return "recetas"
    if any(
        k in full_text
        for k in [
            "salud",
            "enfermedad",
            "sintomas",
            "remedio",
            "beneficios",
            "propiedades",
            "medico",
            "doctor",
            "dolor",
            "curcuma",
            "colageno",
            "vitamina",
            "dieta",
            "organismo",
        ]
    ):
        return "salud_bienestar"
    if any(
        k in full_text
        for k in [
            "partido",
            "futbol",
            "gol",
            "seleccion",
            "champions",
            "mundial",
            "liga",
            "marcador",
            "resultado",
            "pelicula",
            "serie",
            "estreno",
            "actor",
        ]
    ):
        return "deportes_entretenimiento"

    return "general"


def _generar_fallbacks_inteligentes(keyword_base: str, tags_plana: list[str]) -> list[dict]:
    kw_l = keyword_base.lower()
    cat = _detectar_categoria_contextual(keyword_base)

    if cat == "noticias_sismos":
        return [
            {
                "id": "cobertura_vivo",
                "angulo": "Último Reporte & Cobertura en Vivo",
                "icono": "newspaper",
                "titular_h1": f"Último reporte sobre {kw_l}: epicentro, magnitud y municipios donde se sintió",
                "titular_discover": "Sismo en Colombia hoy: informe oficial de magnitud y zonas donde fue percibido",
                "bajada": "El Servicio Sismológico y las autoridades de gestión del riesgo emiten el balance oficial tras el reciente temblor.",
                "interlink_sugerido": "¿Qué hacer tras un sismo en Colombia? Pasos clave para verificar estructuras en casa",
                "tesis_editorial": "Reportar datos oficiales de magnitud, profundidad, mapa de afectación y declaraciones institucionales.",
                "longitud_sugerida": "800 - 1.200 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "prevencion_emergencia",
                "angulo": "Prevención & Protocolos de Emergencia",
                "icono": "warning",
                "titular_h1": "¿Qué debe contener el kit de emergencia ante un sismo y cómo actuar durante el temblor?",
                "titular_discover": "El protocolo de seguridad definitivo ante sismos en Colombia: lo que no puede faltar en su hogar",
                "bajada": "Recomendaciones de los organismos de socorro para proteger a su familia y mascotas ante eventos telúricos.",
                "interlink_sugerido": "¿Cómo preparar a niños y adultos mayores durante una alerta de sismo?",
                "tesis_editorial": "Guía práctica con elementos del maletín de vida, puntos de encuentro y recomendaciones en edificios o vía pública.",
                "longitud_sugerida": "900 - 1.300 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "explicativo_geologico",
                "angulo": "Explicación Científica & Geológica",
                "icono": "public",
                "titular_h1": "¿Por qué tiembla tanto en Colombia y cuáles son las zonas de mayor actividad sísmica?",
                "titular_discover": "La razón científica por la que ocurren constantes temblores en Colombia explicada por geólogos",
                "bajada": "Especialistas analizan la interacción de las placas tectónicas en el territorio nacional.",
                "interlink_sugerido": "Conozca la falla geológica de la Mesa de los Santos y su impacto en la sismicidad del país",
                "tesis_editorial": "Explicar con datos sencillos la tectónica de placas en Colombia y los antecedentes históricos.",
                "longitud_sugerida": "1.000 - 1.400 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "lineas_atencion",
                "angulo": "Líneas de Atención & Reporte de Daños",
                "icono": "phone_in_talk",
                "titular_h1": "Números de emergencia y canales para reportar afectaciones tras sismo en Colombia",
                "titular_discover": "¿Dónde reportar grietas o emergencias tras el sismo? Lista completa de canales oficiales",
                "bajada": "Directorio nacional de Cruz Roja, Defensa Civil, Bomberos y consejos para evaluar daños en edificaciones.",
                "interlink_sugerido": "Cómo solicitar una inspección técnica de riesgo en su municipio",
                "tesis_editorial": "Directorio completo de contactos de emergencia, recomendaciones de evaluación visual de daños y números de socorro.",
                "longitud_sugerida": "700 - 1.000 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "mitos_alarmas",
                "angulo": "Verificación de Datos & Falsas Alarmas",
                "icono": "fact_check",
                "titular_h1": "¿Se pueden predecir los sismos? Mitos y realidades aclarados por expertos sismólogos",
                "titular_discover": "La verdad detrás de las cadenas de redes sobre supuestos terremotos inminentes en Colombia",
                "bajada": "Desmintiendo falsos rumores que circulan en redes sociales sobre predicciones de sismos y luces en el cielo.",
                "interlink_sugerido": "¿Cómo activar las alertas sísmicas tempranas en su teléfono móvil?",
                "tesis_editorial": "Desmentir cadenas falsas, explicar por qué la ciencia aún no predice sismos y enseñar a verificar fuentes oficiales.",
                "longitud_sugerida": "900 - 1.200 palabras",
                "tags_recomendados": tags_plana[:6],
            },
        ]
    elif cat == "finanzas_banca":
        return [
            {
                "id": "guia_solicitud",
                "angulo": "Guía de Solicitud & Requisitos",
                "icono": "credit_card",
                "titular_h1": f"Requisitos y paso a paso para solicitar {kw_l} en Colombia sin contratiempos",
                "titular_discover": f"¿Cómo tramitar {kw_l}? Documentos, ingresos mínimos y aprobación rápida",
                "bajada": "Guía detallada con los criterios de evaluación, canales digitales de solicitud y tiempo de entrega.",
                "interlink_sugerido": f"¿Qué puntaje crediticio se requiere en Datacrédito para calificar a {kw_l}?",
                "tesis_editorial": "Explicar el procedimiento de estudio crediticio, ingresos requeridos y consejos para elevar la probabilidad de aprobación.",
                "longitud_sugerida": "900 - 1.300 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "comparativa_tasas",
                "angulo": "Comparativa de Tasas & Costos",
                "icono": "calculate",
                "titular_h1": f"Tasas de interés y costos en {kw_l}: ¿conviene frente a otras opciones?",
                "titular_discover": f"El análisis completo de costos y tasa de interés efectiva anual de {kw_l}",
                "bajada": "Analizamos comisiones, cuotas de manejo y cobros adicionales para que tome la mejor decisión financiera.",
                "interlink_sugerido": "Comparativa entre las principales opciones bancarias en Colombia",
                "tesis_editorial": "Desglosar tasa EA, seguro de deudores, cuota de manejo y cálculo de compras a 1 cuota sin intereses.",
                "longitud_sugerida": "1.000 - 1.400 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "ahorro_cuotas",
                "angulo": "Estrategias de Ahorro & Exoneración",
                "icono": "savings",
                "titular_h1": f"¿Cómo no pagar cuota de manejo en {kw_l}? El truco legal que pocos usuarios usan",
                "titular_discover": f"El secreto para exonerar la cuota de manejo en {kw_l}",
                "bajada": "Especialistas en finanzas personales revelan las metas de consumo y convenios que eliminan cobros innecesarios.",
                "interlink_sugerido": "Cómo solicitar la unificación de deudas o rebaja de tasa con su entidad bancaria",
                "tesis_editorial": "Tácticas comprobadas de negociación bancaria, cashback y metas mensuales para exonerar cuotas de manejo.",
                "longitud_sugerida": "850 - 1.200 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "seguridad_antifraude",
                "angulo": "Seguridad & Prevención de Fraudes",
                "icono": "shield",
                "titular_h1": f"¿Cómo proteger {kw_l} contra clonación y estafas virtuales en Colombia?",
                "titular_discover": f"Las recomendaciones de ciberseguridad para evitar robos y cargos no reconocidos en {kw_l}",
                "bajada": "Claves para configurar la clave dinámica, límites transaccionales y bloquear operaciones sospechosas en tiempo real.",
                "interlink_sugerido": "¿Qué hacer inmediatamente si le cobran una transacción no autorizada?",
                "tesis_editorial": "Paso a paso para bloquear tarjetas digitales, apagar compras internacionales y presentar reclamos de reversión de pago.",
                "longitud_sugerida": "900 - 1.200 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "preguntas_frecuentes",
                "angulo": "Preguntas Frecuentes & Mitos Financieros",
                "icono": "help",
                "titular_h1": f"Todo lo que debe saber sobre {kw_l} antes de firmar o contratar",
                "titular_discover": f"Los datos cruciales sobre {kw_l} que la mayoría de clientes descubre demasiado tarde",
                "bajada": "Resolvemos las dudas más recurrentes sobre fechas de corte, pago mínimo e impacto en su historial crediticio.",
                "interlink_sugerido": f"¿Afecta su score financiero pagar el pago mínimo en {kw_l}?",
                "tesis_editorial": "Aclarar mitos de tarjetas de crédito, fechas de pago vs fechas de corte y cálculo de intereses corrientes y de mora.",
                "longitud_sugerida": "1.000 - 1.300 palabras",
                "tags_recomendados": tags_plana[:6],
            },
        ]
    else:
        # Fallback general adaptativo
        return [
            {
                "id": "guia_completa",
                "angulo": "Guía Completa & Explicación",
                "icono": "auto_stories",
                "titular_h1": f"Todo lo que necesita saber sobre {kw_l}: guía práctica y conceptos clave",
                "titular_discover": f"La guía definitiva sobre {kw_l}: detalles, respuestas y lo que dicen los expertos",
                "bajada": "Un análisis exhaustivo basado en las preguntas y búsquedas más frecuentes de los usuarios.",
                "interlink_sugerido": f"Aspectos clave y recomendaciones esenciales a tener en cuenta sobre {kw_l}",
                "tesis_editorial": f"Explicar con claridad qué es {kw_l}, cómo funciona, beneficios y recomendaciones de uso.",
                "longitud_sugerida": "900 - 1.300 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "beneficios_claves",
                "angulo": "Beneficios & Recomendaciones",
                "icono": "star",
                "titular_h1": f"Los beneficios principales de {kw_l} y cómo aprovecharlos al máximo",
                "titular_discover": f"Por qué {kw_l} está ganando interés y cuáles son sus mayores ventajas",
                "bajada": "Analizamos los aspectos más destacados y las recomendaciones prácticas respaldadas por especialistas.",
                "interlink_sugerido": f"Errores habituales al utilizar {kw_l} y cómo evitarlos",
                "tesis_editorial": "Desglosar las ventajas más valoradas por los usuarios con consejos concretos para potenciar sus resultados.",
                "longitud_sugerida": "900 - 1.200 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "paso_a_paso",
                "angulo": "Paso a Paso & Procedimiento",
                "icono": "format_list_numbered",
                "titular_h1": f"Cómo aplicar u optimizar {kw_l}: procedimiento paso a paso",
                "titular_discover": f"El método sencillo y efectivo para utilizar {kw_l} sin complicaciones",
                "bajada": "Instrucciones claras y detalladas numeradas para obtener los mejores resultados en poco tiempo.",
                "interlink_sugerido": f"Herramientas y recursos complementarios para optimizar {kw_l}",
                "tesis_editorial": "Proporcionar un tutorial paso a paso con recomendaciones técnicas o prácticas de fácil ejecución.",
                "longitud_sugerida": "800 - 1.100 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "comparativa_analisis",
                "angulo": "Comparativa & Alternativas",
                "icono": "compare_arrows",
                "titular_h1": f"¿Conviene {kw_l}? Análisis frente a las principales alternativas del mercado",
                "titular_discover": f"La comparativa sobre {kw_l}: ventajas, desventajas y recomendaciones",
                "bajada": "Evaluamos las características clave y alternativas para ayudarte a tomar la decisión correcta.",
                "interlink_sugerido": "Factores determinantes al elegir la mejor opción según sus necesidades",
                "tesis_editorial": "Comparar puntos fuertes, puntos débiles y relación calidad/precio frente a otras opciones similares.",
                "longitud_sugerida": "1.000 - 1.400 palabras",
                "tags_recomendados": tags_plana[:6],
            },
            {
                "id": "mitos_realidades",
                "angulo": "Mitos vs Realidades",
                "icono": "psychology_alt",
                "titular_h1": f"Mitos y verdades sobre {kw_l}: desmintiendo las creencias más populares",
                "titular_discover": f"Lo que nadie le dice sobre {kw_l}: la verdad respaldada por datos",
                "bajada": "Separamos la ficción de los hechos comprobados para aclarar las dudas más comunes.",
                "interlink_sugerido": f"Preguntas frecuentes que todo usuario se hace sobre {kw_l}",
                "tesis_editorial": "Analizar los mitos más difundidos en redes o internet y contraponerlos con evidencia y opinión de expertos.",
                "longitud_sugerida": "950 - 1.300 palabras",
                "tags_recomendados": tags_plana[:6],
            },
        ]


def generar_ideas_notas_angulos(
    keyword_base: str,
    sugerencias: list[str] = None,
    preguntas_paa: list[str] = None,
    pais: str = "Colombia",
    tags_reales: dict = None,
) -> list[dict]:
    """
    Genera 5 ideas de notas clasificadas por ángulos editoriales periodísticos
    adaptados estrictamente al contexto temático real de la keyword.
    """
    sugs = (sugerencias or [])[:15]
    paas = (preguntas_paa or [])[:10]
    tags_info = tags_reales or obtener_tags_reales_google(keyword_base, sugs, paas, pais)
    tags_plana = tags_info.get("tags_lista_plana", [keyword_base])

    cat = _detectar_categoria_contextual(keyword_base, sugs, paas)
    fallback_ideas = _generar_fallbacks_inteligentes(keyword_base, tags_plana)

    if not GROQ_API_KEY:
        return fallback_ideas

    prompt = (
        f"Eres el Editor Jefe de un prestigioso periódico digital en {pais}.\n"
        f"Tu objetivo es crear 5 IDEAS DE NOTAS PERIODÍSTICAS Y DE BLOG extraordinarias, coherentes y de alto CTR sobre el tema: '{keyword_base}'.\n"
        f"CATEGORÍA TEMÁTICA DETECTADA: {cat.upper()}.\n\n"
        "REGLA DE ORO ABSOLUTA:\n"
        "DEBES adaptar los 5 ángulos de manera 100% LÓGICA y REALISTA a la naturaleza temática del término. Está estrictamente PROHIBIDO inventar disparates como 'receta de sismo', 'consumir tarjeta de crédito', 'poner una cuchara para terremoto', 'ahorrar energía con temblor', etc. Todo debe tener sentido periodístico profesional real.\n\n"
        f"Evidencia de búsquedas reales (Google Suggest): {json.dumps(sugs, ensure_ascii=False)}\n"
        f"Preguntas que hace la gente (PAA): {json.dumps(paas, ensure_ascii=False)}\n"
        f"Tags Reales de Google Trends disponibles: {json.dumps(tags_plana, ensure_ascii=False)}\n\n"
        "Genera 5 ángulos totalmente alineados con la temática real (ej. si es sismo/noticia: reporte en vivo, prevención, causas científicas, números de emergencia, mitos sismológicos; si es tarjeta/banco: requisitos, comparativa de tasas, cuotas de manejo, seguridad antifraude, preguntas frecuentes; etc.).\n\n"
        "Devuelve ÚNICAMENTE un array JSON con 5 objetos estructurados como sigue:\n"
        "[\n"
        "  {\n"
        '    "id": "angulo_1",\n'
        '    "angulo": "Nombre del Ángulo Coherente",\n'
        '    "icono": "newspaper",\n'
        '    "titular_h1": "Titular H1 periodístico con alto CTR",\n'
        '    "titular_discover": "Titular magnético para Google Discover",\n'
        '    "bajada": "Resumen en dos líneas del enfoque...",\n'
        '    "interlink_sugerido": "¿Pregunta o tema conectado para enlazar?",\n'
        '    "tesis_editorial": "Qué explicará el artículo...",\n'
        '    "longitud_sugerida": "900 - 1.200 palabras",\n'
        '    "tags_recomendados": ["tag1", "tag2", "tag3"]\n'
        "  }\n"
        "]"
    )

    try:
        data = _post_groq_json(prompt, timeout=45)
        if isinstance(data, list) and len(data) >= 3:
            return data
    except Exception as e:
        logger.warning("Error generando ideas de notas con Groq: %s", e)

    return fallback_ideas


def redactar_nota_editorial(
    keyword_base: str,
    angulo: str = "Trucos y Hacks Cotidianos",
    titular_h1: str = None,
    sugerencias: list[str] = None,
    preguntas_paa: list[str] = None,
    pais: str = "Colombia",
    tags_reales: dict = None,
) -> dict:
    """
    Redacta una nota periodística / artículo SEO completo siguiendo estrictamente
    la fórmula y estilo de 'Referencias de estilo de redacción.docx'.
    """
    sugs = (sugerencias or [])[:15]
    paas = (preguntas_paa or [])[:10]
    tags_info = tags_reales or obtener_tags_reales_google(keyword_base, sugs, paas, pais)
    tags_plana = tags_info.get("tags_lista_plana", [keyword_base])
    tags_clasificados = tags_info.get("tags_clasificados", {})

    h1_definido = titular_h1 or f"¿Para qué sirve {keyword_base.lower()} y por qué recomiendan hacerlo?"

    fallback_nota = {
        "titular_h1": h1_definido,
        "titular_discover": f"El método con {keyword_base.lower()} que está transformando hogares",
        "bajada": f"Especialistas y fabricantes detallan las claves principales sobre {keyword_base.lower()} para aprovechar sus beneficios y evitar errores comunes.",
        "interlink_sugerido": f"¿Tiene dudas sobre {keyword_base.lower()}? Su origen y funcionamiento es clave para evitar inconvenientes",
        "parrafo_intro_1": f"El uso de {keyword_base.lower()} suele cobrar relevancia en situaciones cotidianas donde las personas buscan optimizar sus rutinas, mejorar el bienestar del hogar o solucionar dudas frecuentes que impactan la vida diaria.",
        "parrafo_intro_2": f"Frente a este escenario, diversas recomendaciones de especialistas y experiencias prácticas señalan que prestar atención a los detalles de {keyword_base.lower()} permite obtener resultados efectivos sin incurrir en gastos adicionales ni procedimientos complejos.",
        "respaldo_autoridad": "De acuerdo con análisis técnicos y estudios especializados en el sector, la correcta implementación de estos métodos contribuye significativamente al confort y a la durabilidad de los recursos del hogar.",
        "foto_1": {
            "pie_de_foto": f"El uso adecuado de {keyword_base.lower()} aporta múltiples ventajas cotidianas.",
            "credito": "Foto: Imagen generada por IA",
        },
        "secciones_h2": [
            {
                "h2": f"¿Cómo funciona y cuáles son los fundamentos de {keyword_base.title()}?",
                "parrafo_intro_seccion": "La propuesta de este método consiste en aplicar principios prácticos orientados a obtener la máxima eficiencia:",
                "bullets": [
                    {
                        "negrita": "Principio fundamental:",
                        "texto": "Aprovecha los factores del entorno para canalizar y potenciar el efecto deseado de forma controlada.",
                    },
                    {
                        "negrita": "Facilidad de aplicación:",
                        "texto": "No requiere herramientas especializadas y puede implementarse en pocos minutos en cualquier vivienda.",
                    },
                    {
                        "negrita": "Impacto inmediato:",
                        "texto": "Permite notar cambios favorables desde las primeras horas de uso continuo.",
                    },
                ],
            },
            {
                "h2": "Beneficios clave y recomendaciones prácticas",
                "parrafo_intro_seccion": "Aunque no sustituye soluciones estructurales complejas, su aplicación constante aporta beneficios concretos:",
                "bullets": [
                    {
                        "negrita": "Optimización de recursos:",
                        "texto": "Contribuye a reducir gastos innecesarios y prolongar la vida útil de los elementos involucrados.",
                    },
                    {
                        "negrita": "Prevención de inconvenientes:",
                        "texto": "Disminuye notablemente la probabilidad de que surjan deterioros acumulativos con el paso del tiempo.",
                    },
                    {
                        "negrita": "Mantenimiento preventivo:",
                        "texto": "Se recomienda complementar esta práctica con revisiones periódicas cada tres o seis meses.",
                    },
                ],
            },
        ],
        "foto_2": {
            "pie_de_foto": f"Recomendaciones de los expertos para optimizar {keyword_base.lower()}.",
            "credito": "Foto: iStock",
        },
        "seccion_preguntas_frecuentes": {
            "h2": f"¿Qué hacer si surgen complicaciones con {keyword_base.lower()}?",
            "parrafo": "Para situaciones donde no se obtengan los resultados esperados, se aconseja verificar las condiciones del entorno y consultar a entidades autorizadas o técnicos certificados para un diagnóstico formal.",
        },
        "cierre_responsable": "Antes de adoptar cualquier cambio radical en sus rutinas o en el mantenimiento de su hogar, es aconsejable consultar con profesionales calificados o revisar los manuales oficiales del fabricante. Su incorporación debe entenderse como un complemento accesible y responsable dentro de los hábitos generales.",
        "tags_reales": tags_plana[:10],
        "tags_clasificados": tags_clasificados,
    }

    if not GROQ_API_KEY:
        return fallback_nota

    cat = _detectar_categoria_contextual(keyword_base, sugs, paas)
    prompt = (
        f"Eres un Redactor Periodístico Senior y Estratega SEO de un prestigioso medio digital en {pais}.\n"
        f"Debes redactar una NOTA PERIODÍSTICA / ARTÍCULO COMPLETO sobre: '{keyword_base}'.\n"
        f"Ángulo editorial: '{angulo}'.\n"
        f"Titular base propuesto: '{h1_definido}'.\n"
        f"CATEGORÍA TEMÁTICA DETECTADA: {cat.upper()}.\n\n"
        "ESTILO Y ESTRUCTURA OBLIGATORIA (Basada en el Manual de Estilo Periodístico de Referencia):\n"
        "1. titular_h1: Titular periodístico de alto impacto, curiosidad o beneficio adaptado 100% al tema real (sin mezclar trucos de cocina ni hogar si el tema es un desastre, banco o tecnología).\n"
        "2. titular_discover: Titular alternativo para Google Discover / Redes Sociales.\n"
        "3. bajada: 1 a 2 líneas concisas que resuman la tesis de la nota.\n"
        "4. interlink_sugerido: Pregunta o tema directamente relacionado para enlazar internamente.\n"
        "5. parrafo_intro_1 y parrafo_intro_2: Introducción que conecta de forma profesional con la situación real del lector.\n"
        "6. respaldo_autoridad: Cita o mención de entidades oficiales, estudios o expertos calificados del sector.\n"
        "7. foto_1 y foto_2: Pie de foto contextualizado y crédito ('Foto: iStock' o 'Foto: Imagen generada por IA').\n"
        "8. secciones_h2: Array de 2 a 3 secciones H2. Cada H2 debe tener un párrafo introductorio y un array 'bullets' donde CADA bullet tiene 'negrita' (nombre del concepto con dos puntos) y 'texto' (explicación concisa y accionable).\n"
        "9. seccion_preguntas_frecuentes: Sección H2 resolviendo una duda crucial sobre el tema real.\n"
        "10. cierre_responsable: Párrafo final con advertencia técnica/médica/oficial, llamado a la moderación o recomendación de consultar expertos.\n"
        f"11. tags_reales: Array de 6 a 10 tags seleccionados EXCLUSIVAMENTE de esta lista de Google Trends y SERP: {json.dumps(tags_plana, ensure_ascii=False)}. ¡PROHIBIDO INVENTAR ETIQUETAS!\n\n"
        "Devuelve ÚNICAMENTE un objeto JSON válido con esta estructura:\n"
        "{\n"
        '  "titular_h1": "...",\n'
        '  "titular_discover": "...",\n'
        '  "bajada": "...",\n'
        '  "interlink_sugerido": "...",\n'
        '  "parrafo_intro_1": "...",\n'
        '  "parrafo_intro_2": "...",\n'
        '  "respaldo_autoridad": "...",\n'
        '  "foto_1": {"pie_de_foto": "...", "credito": "Foto: Imagen generada por IA"},\n'
        '  "secciones_h2": [\n'
        "    {\n"
        '      "h2": "...",\n'
        '      "parrafo_intro_seccion": "...",\n'
        '      "bullets": [\n'
        '        {"negrita": "Concepto clave:", "texto": "Explicación detallada..."},\n'
        '        {"negrita": "Otro punto:", "texto": "Explicación..."}\n'
        "      ]\n"
        "    }\n"
        "  ],\n"
        '  "foto_2": {"pie_de_foto": "...", "credito": "Foto: iStock"},\n'
        '  "seccion_preguntas_frecuentes": {"h2": "...", "parrafo": "..."},\n'
        '  "cierre_responsable": "...",\n'
        '  "tags_reales": ["tag1", "tag2", "tag3"]\n'
        "}"
    )

    try:
        data = _post_groq_json(prompt, timeout=60)
        if isinstance(data, dict) and "titular_h1" in data and "secciones_h2" in data:
            data["tags_clasificados"] = tags_clasificados
            # Asegurar que tags_reales no esté vacío
            if not data.get("tags_reales"):
                data["tags_reales"] = tags_plana[:8]
            return data
    except Exception as e:
        logger.warning("Error redactando nota editorial con Groq: %s", e)

    return fallback_nota


def exportar_nota_markdown(nota_data: dict) -> str:
    """Convierte la nota generada en formato Markdown estructurado."""
    h1 = nota_data.get("titular_h1", "")
    bajada = nota_data.get("bajada", "")
    interlink = nota_data.get("interlink_sugerido", "")
    p1 = nota_data.get("parrafo_intro_1", "")
    p2 = nota_data.get("parrafo_intro_2", "")
    autoridad = nota_data.get("respaldo_autoridad", "")
    f1 = nota_data.get("foto_1", {})
    f2 = nota_data.get("foto_2", {})
    cierre = nota_data.get("cierre_responsable", "")
    tags = nota_data.get("tags_reales", [])

    lines = []
    lines.append(f"# {h1}\n")
    if bajada:
        lines.append(f"> **{bajada}**\n")
    if interlink:
        lines.append(f"🔗 *Le puede interesar:* [{interlink}](#)\n")

    if p1:
        lines.append(f"{p1}\n")
    if p2:
        lines.append(f"{p2}\n")
    if autoridad:
        lines.append(f"{autoridad}\n")

    if f1 and f1.get("pie_de_foto"):
        lines.append(f"📷 *{f1.get('pie_de_foto')}* — `{f1.get('credito', 'Foto: iStock')}`\n")

    for sec in nota_data.get("secciones_h2", []):
        lines.append(f"\n## {sec.get('h2')}\n")
        if sec.get("parrafo_intro_seccion"):
            lines.append(f"{sec.get('parrafo_intro_seccion')}\n")
        for b in sec.get("bullets", []):
            negrita = b.get("negrita", "")
            texto = b.get("texto", "")
            lines.append(f"- **{negrita}** {texto}")

    if f2 and f2.get("pie_de_foto"):
        lines.append(f"\n📷 *{f2.get('pie_de_foto')}* — `{f2.get('credito', 'Foto: iStock')}`\n")

    faq = nota_data.get("seccion_preguntas_frecuentes", {})
    if faq and faq.get("h2"):
        lines.append(f"\n## {faq.get('h2')}\n")
        lines.append(f"{faq.get('parrafo', '')}\n")

    if cierre:
        lines.append(f"\n{cierre}\n")

    if tags:
        lines.append("\n---")
        lines.append(f"🏷️ **Tags de Google Trends (Reales):** `{'`, `'.join(tags)}`")

    return "\n".join(lines)


def exportar_nota_html(nota_data: dict) -> str:
    """Convierte la nota generada en HTML semántico listo para copiar a WordPress / CMS."""
    h1 = nota_data.get("titular_h1", "")
    bajada = nota_data.get("bajada", "")
    interlink = nota_data.get("interlink_sugerido", "")
    p1 = nota_data.get("parrafo_intro_1", "")
    p2 = nota_data.get("parrafo_intro_2", "")
    autoridad = nota_data.get("respaldo_autoridad", "")
    f1 = nota_data.get("foto_1", {})
    f2 = nota_data.get("foto_2", {})
    cierre = nota_data.get("cierre_responsable", "")

    html = []
    html.append(f"<h1>{h1}</h1>")
    if bajada:
        html.append(f'<p class="lead"><strong>{bajada}</strong></p>')
    if interlink:
        html.append(f'<p class="interlink"><em>Lea también:</em> <a href="#">{interlink}</a></p>')

    if p1:
        html.append(f"<p>{p1}</p>")
    if p2:
        html.append(f"<p>{p2}</p>")
    if autoridad:
        html.append(f"<p>{autoridad}</p>")

    if f1 and f1.get("pie_de_foto"):
        html.append(
            f"<figure><figcaption>{f1.get('pie_de_foto')} <em>({f1.get('credito', 'Foto: iStock')})</em></figcaption></figure>"
        )

    for sec in nota_data.get("secciones_h2", []):
        html.append(f"<h2>{sec.get('h2')}</h2>")
        if sec.get("parrafo_intro_seccion"):
            html.append(f"<p>{sec.get('parrafo_intro_seccion')}</p>")
        html.append("<ul>")
        for b in sec.get("bullets", []):
            html.append(f"<li><strong>{b.get('negrita', '')}</strong> {b.get('texto', '')}</li>")
        html.append("</ul>")

    if f2 and f2.get("pie_de_foto"):
        html.append(
            f"<figure><figcaption>{f2.get('pie_de_foto')} <em>({f2.get('credito', 'Foto: iStock')})</em></figcaption></figure>"
        )

    faq = nota_data.get("seccion_preguntas_frecuentes", {})
    if faq and faq.get("h2"):
        html.append(f"<h2>{faq.get('h2')}</h2>")
        html.append(f"<p>{faq.get('parrafo', '')}</p>")

    if cierre:
        html.append(f"<p>{cierre}</p>")

    return "\n".join(html)


def exportar_nota_docx(nota_data: dict, filepath: str = None) -> io.BytesIO:
    """
    Genera un archivo Word (.docx) con formato editorial impecable:
    - Tipografía Arial limpia
    - Título H1 estilizado
    - Bajada en caja destacada
    - Interlink sugerido
    - Encabezados H2 coloreados
    - Viñetas con negrita
    - Cuadros de foto sugerida con pie de página
    - Tabla de Tags Reales de Google Trends
    """
    import docx
    from docx.shared import Inches, Pt, RGBColor

    doc = docx.Document()

    # Configuración de márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. Titular H1
    h1_text = nota_data.get("titular_h1", "Nota Periodística")
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run(h1_text)
    run_h1.font.name = "Arial"
    run_h1.font.size = Pt(20)
    run_h1.font.bold = True
    run_h1.font.color.rgb = RGBColor(13, 28, 50)  # Primary dark

    # 2. Bajada / Subtítulo
    bajada_text = nota_data.get("bajada", "")
    if bajada_text:
        p_bajada = doc.add_paragraph()
        run_bajada = p_bajada.add_run(bajada_text)
        run_bajada.font.name = "Arial"
        run_bajada.font.size = Pt(12)
        run_bajada.font.bold = True
        run_bajada.font.color.rgb = RGBColor(70, 80, 95)
        p_bajada.paragraph_format.space_after = Pt(14)

    # 3. Interlink sugerido
    interlink_text = nota_data.get("interlink_sugerido", "")
    if interlink_text:
        p_inter = doc.add_paragraph()
        r_lbl = p_inter.add_run("🔗 Tema Relacionado (Interlink): ")
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(0, 103, 127)  # Primary cyan/blue
        r_lnk = p_inter.add_run(interlink_text)
        r_lnk.font.italic = True
        p_inter.paragraph_format.space_after = Pt(14)

    # 4. Párrafos introductorios
    p1_text = nota_data.get("parrafo_intro_1", "")
    if p1_text:
        p1 = doc.add_paragraph(p1_text)
        p1.paragraph_format.space_after = Pt(10)

    p2_text = nota_data.get("parrafo_intro_2", "")
    if p2_text:
        p2 = doc.add_paragraph(p2_text)
        p2.paragraph_format.space_after = Pt(10)

    aut_text = nota_data.get("respaldo_autoridad", "")
    if aut_text:
        p_aut = doc.add_paragraph(aut_text)
        p_aut.paragraph_format.space_after = Pt(14)

    # 5. Foto 1
    f1 = nota_data.get("foto_1", {})
    if f1 and f1.get("pie_de_foto"):
        p_f1 = doc.add_paragraph()
        r_f1_t = p_f1.add_run(f"📷 [Imagen Sugerida] {f1.get('pie_de_foto')}")
        r_f1_t.font.italic = True
        r_f1_c = p_f1.add_run(f"  {f1.get('credito', 'Foto: Imagen generada por IA')}")
        r_f1_c.font.bold = True
        r_f1_c.font.size = Pt(9.5)
        p_f1.paragraph_format.space_after = Pt(16)

    # 6. Secciones H2 y Bullets con negrita
    for sec in nota_data.get("secciones_h2", []):
        h2 = doc.add_heading(level=2)
        r_h2 = h2.add_run(sec.get("h2", ""))
        r_h2.font.name = "Arial"
        r_h2.font.size = Pt(14)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(0, 103, 127)

        if sec.get("parrafo_intro_seccion"):
            p_sec_intro = doc.add_paragraph(sec.get("parrafo_intro_seccion"))
            p_sec_intro.paragraph_format.space_after = Pt(8)

        for b in sec.get("bullets", []):
            p_bullet = doc.add_paragraph(style="List Bullet")
            r_negrita = p_bullet.add_run(f"{b.get('negrita', '')} ")
            r_negrita.font.bold = True
            p_bullet.add_run(b.get("texto", ""))
            p_bullet.paragraph_format.space_after = Pt(4)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 7. Foto 2
    f2 = nota_data.get("foto_2", {})
    if f2 and f2.get("pie_de_foto"):
        p_f2 = doc.add_paragraph()
        r_f2_t = p_f2.add_run(f"📷 [Imagen Sugerida] {f2.get('pie_de_foto')}")
        r_f2_t.font.italic = True
        r_f2_c = p_f2.add_run(f"  {f2.get('credito', 'Foto: iStock')}")
        r_f2_c.font.bold = True
        r_f2_c.font.size = Pt(9.5)
        p_f2.paragraph_format.space_after = Pt(16)

    # 8. Preguntas frecuentes / Sección final
    faq = nota_data.get("seccion_preguntas_frecuentes", {})
    if faq and faq.get("h2"):
        h2_faq = doc.add_heading(level=2)
        r_faq = h2_faq.add_run(faq.get("h2", ""))
        r_faq.font.name = "Arial"
        r_faq.font.size = Pt(14)
        r_faq.font.bold = True
        r_faq.font.color.rgb = RGBColor(0, 103, 127)

        p_faq = doc.add_paragraph(faq.get("parrafo", ""))
        p_faq.paragraph_format.space_after = Pt(14)

    # 9. Cierre responsable
    cierre_text = nota_data.get("cierre_responsable", "")
    if cierre_text:
        p_cierre = doc.add_paragraph(cierre_text)
        p_cierre.paragraph_format.space_after = Pt(20)

    # 10. Bloque de Tags Reales de Google Trends
    tags = nota_data.get("tags_reales", [])
    if tags:
        doc.add_heading("🏷️ Taxonomía y Tags Reales (Google Trends)", level=3)
        p_tags = doc.add_paragraph()
        p_tags.add_run("Lista para CMS (WordPress / Arc): ").font.bold = True
        p_tags.add_run(", ".join(tags))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    if filepath:
        with open(filepath, "wb") as f:
            f.write(buffer.getvalue())
        buffer.seek(0)

    return buffer
